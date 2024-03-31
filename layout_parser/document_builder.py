import json
import math
import os.path
import random
import string
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import cv2
import numpy as np
import reportlab.rl_config
from decouple import config
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from loguru import logger
from reportlab.lib.colors import black
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas
from scipy.cluster.hierarchy import fclusterdata

reportlab.rl_config.warnOnMissingFontGlyphs = 1
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import docx

try:
    from .elements import Block, Document, Line, Page, Word
    from .geometry import estimate_page_angle, resolve_enclosing_bbox, resolve_enclosing_rbbox, rotate_boxes, \
        is_overlap, is_box_contained
except:
    from layout_parser.elements import Block, Document, Line, Page, Word
    from layout_parser.geometry import estimate_page_angle, resolve_enclosing_bbox, resolve_enclosing_rbbox, \
        rotate_boxes, is_overlap, is_box_contained

BASE_DIR = Path(__file__).resolve().parent.parent

fontname = "Roboto"
pdfmetrics.registerFont(
    TTFont(fontname, BASE_DIR / 'data_db/fonts/Roboto-Regular.ttf'))

__all__ = ["DocumentBuilder", "export_docs", "export_pdf_searchable"]


class DocumentBuilder(object):
    """Implements a document builder

    Args:
        resolve_lines: whether words should be automatically grouped into lines
        resolve_blocks: whether lines should be automatically grouped into blocks
        paragraph_break: relative length of the minimum space separating paragraphs
        export_as_straight_boxes: if True, force straight boxes in the export (fit a rectangle
            box to all rotated boxes). Else, keep the boxes format unchanged, no matter what it is.
    """

    def __init__(
            self,
            resolve_lines: bool = True,
            resolve_blocks: bool = True,
            paragraph_break: float = 0.035,
            export_as_straight_boxes: bool = False,
    ) -> None:
        self.resolve_lines = resolve_lines
        self.resolve_blocks = resolve_blocks
        self.paragraph_break = paragraph_break
        self.export_as_straight_boxes = export_as_straight_boxes

    @staticmethod
    def _sort_boxes(boxes: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
        """Sort bounding boxes from top to bottom, left to right

        Args:
            boxes: bounding boxes of shape (N, 4) or (N, 4, 2) (in case of rotated bbox)

        Returns:
            tuple: indices of ordered boxes of shape (N,), boxes
                If straight boxes are passed tpo the function, boxes are unchanged
                else: boxes returned are straight boxes fitted to the straightened rotated boxes
                so that we fit the lines afterwards to the straigthened page
        """
        if boxes.ndim == 3:
            boxes = rotate_boxes(
                loc_preds=boxes,
                angle=-estimate_page_angle(boxes),
                orig_shape=(1024, 1024),
                min_angle=5.0,
            )
            boxes = np.concatenate((boxes.min(axis=1), boxes.max(axis=1)), -1)
        return np.lexsort((boxes[:, 0], boxes[:, 1])), boxes
        # return (boxes[:, 0] + 2 * boxes[:, 3] / np.median(boxes[:, 3] - boxes[:, 1])).argsort(), boxes

    def _resolve_sub_lines(self, boxes: np.ndarray, word_idcs: List[int], width: int, height: int) -> List[List[int]]:
        """Split a line in sub_lines

        Args:
            boxes: bounding boxes of shape (N, 4)
            word_idcs: list of indexes for the words of the line

        Returns:
            A list of (sub-)lines computed from the original line (words)
        """
        lines = []
        # Sort words horizontally
        word_idcs = [word_idcs[idx] for idx in boxes[word_idcs, 0].argsort().tolist()]

        space_widths = []
        for i, idx in enumerate(word_idcs[1:], start=1):
            horiz_break = True

            prev_box = boxes[word_idcs[i - 1]]
            # Compute distance between boxes
            dist = abs(boxes[idx, 0] - prev_box[2])
            if dist < 0:
                continue
            space_widths.append(dist)

        space_widths = sorted(space_widths)
        space_med = np.mean(space_widths[: 5])

        space_widths = []

        # Eventually split line horizontally
        if len(word_idcs) < 2:
            lines.append(word_idcs)
        else:
            sub_line = [word_idcs[0]]
            for i in word_idcs[1:]:
                horiz_break = True

                prev_box = boxes[sub_line[-1]]
                # Compute distance between boxes
                dist = boxes[i, 0] - prev_box[2]

                space_widths.append(dist)

                # If distance between boxes is lower than paragraph break, same sub-line
                if dist < self.paragraph_break * width:
                    horiz_break = False

                # if not horiz_break and len(space_widths) >= 2:
                #     mean_space = sum(space_widths) / len(space_widths)
                #     if dist >= mean_space * 4:
                #         horiz_break = True
                # if not horiz_break and dist >= space_med * 4:
                #     horiz_break = True

                if horiz_break:
                    lines.append(sub_line)
                    sub_line = []
                    space_widths.clear()

                sub_line.append(i)
            lines.append(sub_line)

        return lines

    def _resolve_lines(self, boxes: np.ndarray, width: int, height: int) -> Tuple[List[List[int]], Dict[Any, str]]:
        """Order boxes to group them in lines

        Args:
            boxes: bounding boxes of shape (N, 4) or (N, 4, 2) in case of rotated bbox

        Returns:
            nested list of box indices
        """

        # Sort boxes, and straighten the boxes if they are rotated
        idxs, boxes = self._sort_boxes(boxes)

        # Compute median for boxes heights
        y_med = np.median(boxes[:, 3] - boxes[:, 1])

        lines = []
        words = [idxs[0]]  # Assign the top-left word to the first line
        # Define a mean y-center for the line
        y_center_sum = boxes[idxs[0]][[1, 3]].mean()

        y_med = abs(boxes[0, 3] - boxes[0, 1])
        h_sum = y_med
        soft_add_lines = []
        for idx in idxs[1:]:
            vert_break = True

            # Compute y_dist
            y_dist = abs(boxes[idx][[1, 3]].mean() - y_center_sum / len(words))
            h_w = abs(boxes[idx, 3] - boxes[idx, 1])
            h_avg = h_sum / len(words)
            if 2 > h_w / h_avg > 0.5:
                # If y-center of the box is close enough to mean y-center of the line, same line
                if y_dist < y_med / 2:
                    vert_break = False

                    # Nếu y-center of the box is hơi xa to mean y-center of the line, thì cần check thêm box đó có gần line khác không
                    if y_dist > y_med / 4:
                        soft_add_lines.append(idx)

            if vert_break:
                lines.append(words)
                words = []
                y_center_sum = 0
                h_sum = 0
            if len(words) == 0:
                y_med = abs(boxes[idx, 3] - boxes[idx, 1])
                tem_y_center = boxes[idx][[1, 3]].mean()
                # check lại các box trong soft line add nếu thoả mãn điều kiện thì chuyển line trong soft line xuống dòng dưới
                for word_check in soft_add_lines:
                    y_dist_test = abs(boxes[word_check][[1, 3]].mean() - tem_y_center)
                    if y_dist_test < y_med / 4:
                        words.append(word_check)
                        lines[-1].remove(word_check)
                        y_center_sum += boxes[word_check][[1, 3]].mean()
                        h_sum += abs(boxes[word_check][3] - boxes[word_check][1])
                soft_add_lines.clear()

            words.append(idx)
            y_center_sum += boxes[idx][[1, 3]].mean()
            h_sum += h_w

        # Use the remaining words to form the last(s) line(s)
        if len(words) > 0:
            # Compute sub-lines (horizontal split)
            lines.append(words)
        lines_all = []
        map_words_line_ids = {}
        # Compute sub-lines (horizontal split)
        for words in lines:
            # create line id for in the feature we use this for get full line
            line_id = id_generator()

            sub_lines = self._resolve_sub_lines(boxes, words, width, height)
            for sub_line in sub_lines:
                sub_line_id = id_generator()
                for w in sub_line:
                    map_words_line_ids[w] = f"{line_id}_{sub_line_id}"

            lines_all.extend(sub_lines)
        return lines_all, map_words_line_ids

    @staticmethod
    def _resolve_blocks(boxes: np.ndarray, lines: List[List[int]], width: int, height: int) -> List[List[List[int]]]:
        """Order lines to group them in blocks

        Args:
            boxes: bounding boxes of shape (N, 4) or (N, 4, 2)
            lines: list of lines, each line is a list of idx

        Returns:
            nested list of box indices
        """
        # Resolve enclosing boxes of lines
        if boxes.ndim == 3:
            box_lines: np.ndarray = np.asarray(
                [
                    resolve_enclosing_rbbox([tuple(boxes[idx, :, :]) for idx in line])  # type: ignore[misc]
                    for line in lines
                ]
            )
        else:
            _box_lines = [
                resolve_enclosing_bbox(
                    [(tuple(boxes[idx, :2]), tuple(boxes[idx, 2:])) for idx in line]  # type: ignore[misc]
                )
                for line in lines
            ]
            box_lines = np.asarray([(x1, y1, x2, y2) for ((x1, y1), (x2, y2)) in _box_lines])

        # Compute geometrical features of lines to clusterize
        # Clusterizing only with box centers yield to poor results for complex documents
        if boxes.ndim == 3:
            box_features: np.ndarray = np.stack(
                (
                    (box_lines[:, 0, 0] / width + box_lines[:, 0, 1] / height) / 2,
                    (box_lines[:, 0, 0] / width + box_lines[:, 2, 0] / width) / 2,
                    (box_lines[:, 0, 0] / width + box_lines[:, 2, 1] / height) / 2,
                    (box_lines[:, 0, 1] / height + box_lines[:, 2, 1] / height) / 2,
                    (box_lines[:, 0, 1] / height + box_lines[:, 2, 0] / width) / 2,
                    (box_lines[:, 2, 0] / width + box_lines[:, 2, 1] / height) / 2,
                ),
                axis=-1,
            )
        else:
            box_features = np.stack(
                (
                    # (box_lines[:, 0] + box_lines[:, 3]) / 2 / width,
                    # (box_lines[:, 1] + box_lines[:, 2]) / 2 / height,
                    (box_lines[:, 0] + box_lines[:, 2]) / 2 / width,
                    (box_lines[:, 1] + box_lines[:, 3]) / 2 / height,
                    box_lines[:, 0] / width,
                    box_lines[:, 1] / height,
                    (box_lines[:, 3] - box_lines[:, 1]) / height
                ),
                axis=-1,
            )
        # Compute clusters
        clusters = fclusterdata(box_features, t=0.01, depth=4, criterion="distance", metric="euclidean")

        _blocks: Dict[int, List[int]] = {}
        # Form clusters
        for line_idx, cluster_idx in enumerate(clusters):
            if cluster_idx in _blocks.keys():
                _blocks[cluster_idx].append(line_idx)
            else:
                _blocks[cluster_idx] = [line_idx]

        # Retrieve word-box level to return a fully nested structure
        blocks = [[lines[idx] for idx in block] for block in _blocks.values()]

        return blocks

    def _build_blocks(self, boxes: np.ndarray, word_preds: List[Tuple[str, float]], width: int, height: int) -> List[
        Block]:
        """Gather independent words in structured blocks

        Args:
            boxes: bounding boxes of all detected words of the page, of shape (N, 5) or (N, 4, 2)
            word_preds: list of all detected words of the page, of shape N

        Returns:
            list of block elements
        """

        if boxes.shape[0] != len(word_preds):
            raise ValueError(f"Incompatible argument lengths: {boxes.shape[0]}, {len(word_preds)}")

        if boxes.shape[0] == 0:
            return []

        # Decide whether we try to form lines
        _boxes = boxes
        map_words_line_ids = {}
        if self.resolve_lines:
            lines, map_words_line_ids = self._resolve_lines(_boxes if _boxes.ndim == 3 else _boxes[:, :4], width,
                                                            height)
            # Decide whether we try to form blocks
            if self.resolve_blocks and len(lines) > 1:
                _blocks = self._resolve_blocks(_boxes if _boxes.ndim == 3 else _boxes[:, :4], lines, width, height)
            else:
                _blocks = [lines]
        else:
            # Sort bounding boxes, one line for all boxes, one block for the line
            lines = [self._sort_boxes(_boxes if _boxes.ndim == 3 else _boxes[:, :4])[0]]  # type: ignore[list-item]
            _blocks = [lines]

        # Recovery lines in blocks
        if self.resolve_lines:
            for idx, b in enumerate(_blocks):
                map_old_index = {}
                b_boxes = []
                for l in b:
                    for w in l:
                        map_old_index[len(b_boxes)] = w
                        b_boxes.append(_boxes[w])
                b_boxes = np.array(b_boxes)
                b_lines, b_map_words_line_ids = self._resolve_lines(b_boxes if b_boxes.ndim == 3 else b_boxes[:, :4],
                                                                    width, height)
                b_lines_re = []
                for l in b_lines:
                    wl_re = []
                    for w in l:
                        wl_re.append(map_old_index[w])
                    b_lines_re.append(wl_re)
                _blocks[idx] = b_lines_re

        blocks = [
            Block(
                [
                    Line(
                        [
                            Word(
                                *word_preds[idx],
                                tuple([tuple(pt) for pt in boxes[idx].tolist()]),  # type: ignore[arg-type]
                                map_words_line_ids.get(idx), id_generator()
                            )
                            if boxes.ndim == 3
                            else Word(
                                *word_preds[idx],
                                ((int(boxes[idx, 0]), int(boxes[idx, 1])), (int(boxes[idx, 2]), int(boxes[idx, 3]))),
                                map_words_line_ids.get(idx), id_generator()
                            )
                            for idx in line
                        ]
                    )
                    for line in lines
                ]
            )
            for lines in _blocks
        ]

        return blocks

    def extra_repr(self) -> str:
        return (
            f"resolve_lines={self.resolve_lines}, resolve_blocks={self.resolve_blocks}, "
            f"paragraph_break={self.paragraph_break}, "
            f"export_as_straight_boxes={self.export_as_straight_boxes}"
        )

    def filter_boxes_in_figure(self, boxes, text_preds, figures):
        if figures is None or not any(figures):
            return boxes, text_preds
        output_boxes, output_text_preds = [], []
        for _box, _text_preds, _figures in zip(boxes, text_preds, figures):
            page_box = []
            page_text_preds = []
            if _figures:
                for f_x1, f_y1, f_x2, f_y2 in _figures:
                    for _b, t in zip(_box, _text_preds):
                        x1, y1, x2, y2 = _b
                        if not is_box_contained((f_x1, f_y1, f_x2, f_y2), (x1, y1, x2, y2)):
                            page_box.append(_b)
                            page_text_preds.append(t)
                        else:
                            logger.debug(f"Remove box in figure: {_b} - {t}")
                output_boxes.append(np.array(page_box))
                output_text_preds.append(page_text_preds)
            else:
                output_boxes.append(_box)
                output_text_preds.append(_text_preds)

        return output_boxes, output_text_preds

    def __call__(
            self,
            boxes: List[np.ndarray],
            text_preds: List[List[Tuple[str, float]]],
            page_shapes: List[Tuple[int, int]],
            file_id: str,
            orientations: Optional[List[Dict[str, Any]]] = None,
            languages: Optional[List[Dict[str, Any]]] = None,
            figures: Optional[List[List[Tuple[int, int, int, int]]]] = None,
            tables: Optional[List[List[Tuple[int, int, int, int]]]] = None,
            title_boxes: Optional[List[Tuple[int, int, int, int]]] = None,
            page_numbers: Optional[List[Tuple[int, int, int, int]]] = None,
            signature_boxes: Optional[List[Tuple[int, int, int, int]]] = None
    ) -> Document:
        """Re-arrange detected words into structured blocks

        Args:
            boxes: list of N elements, where each element represents the localization predictions, of shape (*, 5)
                or (*, 6) for all words for a given page
            text_preds: list of N elements, where each element is the list of all word prediction (text + confidence)
            page_shape: shape of each page, of size N

        Returns:
            document object
        """
        if len(boxes) != len(text_preds) or len(boxes) != len(page_shapes):
            raise ValueError("All arguments are expected to be lists of the same size")

        _orientations = (
            orientations if isinstance(orientations, list) else [None] * len(boxes)  # type: ignore[list-item]
        )
        _languages = languages if isinstance(languages, list) else [None] * len(boxes)  # type: ignore[list-item]
        _figures = figures if isinstance(figures, list) else [None] * len(boxes)  # type: ignore[list-item]
        _tables = tables if isinstance(tables, list) else [None] * len(boxes)  # type: ignore[list-item]
        _signature_boxes = signature_boxes if isinstance(signature_boxes, list) else [None] * len(
            boxes)  # type: ignore[list-item]
        _title_boxes = title_boxes if isinstance(title_boxes, list) else [None] * len(boxes)  # type: ignore[list-item]
        _page_numbers = page_numbers if isinstance(page_numbers, list) else [None] * len(
            boxes)  # type: ignore[list-item]

        boxes, text_preds = self.filter_boxes_in_figure(boxes, text_preds, figures)

        if self.export_as_straight_boxes and len(boxes) > 0:
            # If boxes are already straight OK, else fit a bounding rect
            if boxes[0].ndim == 3:
                straight_boxes: List[np.ndarray] = []
                # Iterate over pages
                for p_boxes in boxes:
                    # Iterate over boxes of the pages
                    straight_boxes.append(np.concatenate((p_boxes.min(1), p_boxes.max(1)), 1))
                boxes = straight_boxes

        # BASE_DIR = Path(__file__).resolve().parent.parent

        IMAGE_BASE_DIR = config("IMAGE_BASE_DIR", default="data/image/requests")

        _pages = [
            Page(
                self._build_blocks(
                    page_boxes,
                    word_preds,
                    width=shape[1],
                    height=shape[0]
                ),
                _idx,
                shape,
                orientation,
                language,
                image_path=f"{IMAGE_BASE_DIR}/{file_id}_{_idx}.jpg",
                figures=figure,
                tables=table,
                title_box=title_box,
                page_number=page_number,
                signature_boxes=signature_box
            )
            for
            _idx, shape, page_boxes, word_preds, orientation, language, figure, table, title_box, page_number, signature_box
            in
            zip(
                range(len(boxes)), page_shapes, boxes, text_preds, _orientations, _languages, _figures, _tables,
                _title_boxes, _page_numbers, _signature_boxes
            )
        ]

        return Document(_pages)


def id_generator(size=3, chars=string.ascii_uppercase + string.digits):
    return ''.join(random.choice(chars) for _ in range(size))


FONT_DIFF = 2
DIFF = 20
BASE_FONT_SIZE = 12
POINT2MILI = 0.352778
A4MILI = 210, 297  # w, h


def parse_fontsize(doc: Document):
    font_list = []
    height_dict = {}

    for idx, page in enumerate(doc.pages):
        for block in page.blocks:
            for line in block.lines:
                for word in line.words:
                    (x1, y1), (x2, y2) = word.bbox
                    height = y2 - y1
                    if height in height_dict:
                        height_dict[height] += 1
                    else:
                        height_dict[height] = 1

    c = 0
    mid_height = 0
    for k in sorted(height_dict.keys()):
        if height_dict[k] > c:
            c = height_dict[k]
            mid_height = k

    font = {
        'min_height': 0,
        'max_height': mid_height,
        'font_size': BASE_FONT_SIZE
    }

    font_list.append(font)
    for k in sorted(height_dict.keys()):
        if k > font['max_height']:
            font = {
                'min_height': k,
                'max_height': k + FONT_DIFF,
                'font_size': math.floor(k / mid_height * BASE_FONT_SIZE)
            }
    font_list.append(font)
    return font_list


def get_fontsize(font_list: list, height: int):
    for v in font_list:
        if v["min_height"] <= height <= v["max_height"]:
            return v["font_size"]
    return BASE_FONT_SIZE


def parse_border(doc: Document):
    right = 0
    left = 0
    i = 0
    for idx, page in enumerate(doc.pages):
        for block in page.blocks:
            for line in block.lines:
                for word in line.words:
                    (x1, y1), (x2, y2) = word.bbox
                    if i == 0:
                        left = x1
                    left = min(left, x1)
                    right = max(right, x2)
                    i += 1

    return {
        'left': left + DIFF,
        'right': right - DIFF
    }


def recovery_line_page(document: Document) -> List[List[List[Line]]]:
    pages = []
    for idx, page in enumerate(document.pages):
        lines = {}
        for block in page.blocks:
            for line in block.lines:
                line_id = None
                for word in line.words:
                    if word.parent_line_id:
                        line_id, sub_line_id = word.parent_line_id.split('_')
                        break

                if line_id:
                    if line_id not in lines:
                        lines[line_id] = []
                    lines[line_id].append(line)

        page_lines = []
        if lines:
            for line_id, line in lines.items():
                page_lines.append(line)
        else:
            for idx, page in enumerate(document.pages):
                for block in page.blocks:
                    for line in block.lines:
                        page_lines.append(line)
        pages.append(page_lines)

    return pages


def get_bbox_sub_lines(sub_lines: List[Line]) -> Tuple[int, int, int, int]:
    if len(sub_lines) == 0:
        return 0, 0, 0, 0
    (x1_sl, y1_sl), (x2_sl, y2_sl) = sub_lines[0].bbox
    x1, y1, x2, y2 = x1_sl, y1_sl, x2_sl, y2_sl
    for sub_line in sub_lines:
        (x1_sl, y1_sl), (x2_sl, y2_sl) = sub_line.bbox

        x1 = min(x1, x1_sl)
        y1 = min(y1, y1_sl)
        x2 = max(x2, x2_sl)
        y2 = max(y2, y2_sl)
    return x1, y1, x2, y2


def estimate_space(distance: float, words: List[Word]):
    width_char = []
    for word in words:
        (x1, y1), (x2, y2) = word.bbox
        c_w = (x2 - x1) / len(word.value)
        width_char.append(c_w)
    width_space = sum(width_char) / len(width_char) * 0.5
    return int(distance / width_space)


def pixels_to_inches(pixels, dpi=200):
    inches = pixels / dpi
    return Inches(inches)


def add_text_box(doc, text, left, top, width, height):
    # Create a table with one cell
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    # Set the cell properties to simulate a text box
    cell.left = left
    cell.top = top
    cell.width = width
    cell.height = height

    # Add the text to the cell
    cell.text = text


def export_docs(document: Document, output_docx_path):
    docs_rate_size = 0.4
    logger.info("Export image/pdf to docs")
    doc = docx.Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(6.5)

    font_lists = parse_fontsize(document)
    border = parse_border(document)

    num_pages = len(document.pages)
    page_lines = recovery_line_page(document)
    for idx, page in enumerate(document.pages):
        page_height, page_width = page.dimensions
        lines = page_lines[idx]
        for line in lines:
            x1_l, y1_l, x2_l, y2_l = get_bbox_sub_lines(line)

            p = doc.add_paragraph()
            p.paragraph_format.right_indent = pixels_to_inches(page_width - border['right'])
            p.paragraph_format.left_indent = pixels_to_inches(border['left'])
            p.paragraph_format.line_spacing = pixels_to_inches(y2_l - y1_l + 5)

            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            space_left = abs(x1_l - border['left'])
            space_right = abs(border['right'] - x2_l)
            if space_left > 10 and space_right > 10 and abs(space_left - space_right) < 3:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logger.debug(f"Align CENTER: {line[0].render()}")
            elif space_left < 10 and space_right < 10:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                logger.debug(f"Align JUSTIFY: {line[0].render()}")
            else:
                p.paragraph_format.left_indent = pixels_to_inches(x1_l)
            last_x2 = 0
            for sub_line in line:
                (x1_sl, y1_sl), (x2_sl, y2_sl) = sub_line.bbox
                height_sub_line = y2_sl - y1_sl
                width_sub_line = x2_sl - x1_sl

                font_size = get_fontsize(font_list=font_lists, height=height_sub_line)
                font_size = int(height_sub_line * docs_rate_size)
                run = p.add_run()
                if last_x2 > 0 and x1_sl - last_x2 > 0:
                    n_space = estimate_space(distance=x1_sl - last_x2, words=sub_line.words)
                    tab_space = 100
                    if n_space > tab_space:
                        for _ in range(n_space // tab_space):
                            run.add_tab()
                    else:
                        run.add_text(' ' * n_space)
                    logger.debug(f"Add {n_space} spaces before {sub_line.render()}")

                last_x2 = x2_sl
                run.add_text(sub_line.render())
                run.font.size = Pt(font_size)
                if not "bold":
                    run.bold = True

        if idx < num_pages - 1:
            doc.add_page_break()

    margin = 0
    for section in doc.sections:
        # section.top_margin = Pt(margin)
        # section.bottom_margin = Pt(margin)
        section.left_margin = Pt(margin)
        section.right_margin = Pt(margin)

    doc.save(output_docx_path)


def export_pdf_searchable(document: Document, output_pdf_path):
    pdf_writer = Canvas(output_pdf_path, pagesize=letter, pageCompression=1)
    pdf_writer.setFillColor(black)  # text in black

    temp_dir = tempfile.mkdtemp()
    for page_number, page in enumerate(document.pages):
        # Create a new page with the original image
        image_path = os.path.join(temp_dir, f"{page_number}.jpg")
        cv2.imwrite(image_path, page.image)
        im_h, im_w = page.image.shape[:2]
        ratio_width = letter[0] / im_w
        ratio_height = letter[1] / im_h
        pdf_writer.drawImage(image_path, x=0, y=0, width=letter[0], height=letter[1])

        for block in page.blocks:
            for line in block.lines:
                for word in line.words:
                    text = pdf_writer.beginText(word.bbox[0][0] * ratio_width,
                                                (im_h - word.bbox[0][1]) * ratio_height - 15)
                    text.setTextRenderMode(3)  # Invisible (indicates OCR text)
                    text.setFont(fontname, 13)
                    text.textOut(word.value + ' ')

                    # Add the searchable text to the PDF
                    pdf_writer.drawText(text)

        # If this is not the last page, add a new page
        if page_number < len(document.pages) - 1:
            pdf_writer.showPage()

    pdf_writer.save()


def past(json_path, file_id):
    data = json.load(open(json_path, encoding='utf8'))
    print(data)
    document = Document.from_dict(data)
    document.remove_text_in_figure()
    data_dump = {
        'text_preds': [],
        'boxes': [],
        'orientations': None,
        'page_shapes': [],
        'languages': None,
        'figures': [],
        'tables': [],
        'title_boxes': [],
        'page_numbers': [],
        'file_id': file_id
    }
    for page in document.pages:
        boxes = []
        text_preds = []
        figures = []
        tables = []
        page_shapes = page.dimensions
        for block in page.blocks:
            for line in block.lines:
                for word in line.words:
                    (x1, y1), (x2, y2) = word.bbox
                    boxes.append((x1, y1, x2, y2))
                    text_preds.append((word.value, word.confidence))

        for f in page.figures:
            figures.append(f)

        for f in page.tables:
            tables.append(f)

        data_dump['text_preds'].append(text_preds)
        data_dump['boxes'].append(np.array(boxes))
        data_dump['page_shapes'].append(page_shapes)
        data_dump['figures'].append(figures)
        data_dump['tables'].append(tables)
        data_dump['title_boxes'].append(page.title_box)
        data_dump['page_numbers'].append(page.page_number)

    pickle.dump(data_dump, open('data.pkl', 'wb'))


if __name__ == '__main__':

    import pickle

    case_tests = {
        '1': {
            'json': '9fd85741-0d8e-4660-bc4a-42fa20c60fab.json',
            'imgs': ['9fd85741-0d8e-4660-bc4a-42fa20c60fab_0.jpg']
        },
        '2': {
            'json': 'c34a9446-f68c-42e2-8adc-0ddefde32fc8.json',
            'imgs': ['c34a9446-f68c-42e2-8adc-0ddefde32fc8_0.jpg']
        },
        '3': {
            'json': '8467360a-4aa4-4178-8aaf-ee3bd600af91.json',
            'imgs': ['8467360a-4aa4-4178-8aaf-ee3bd600af91_0.jpg']
        },
        '4': {
            'json': '7672fac2-0104-4ed7-87a2-d7cf769c4831.json',
            'imgs': ['7672fac2-0104-4ed7-87a2-d7cf769c4831_0.jpg']
        },
        '5': {
            'json': '1bea552d-5295-4ee1-a43c-16b527a59859.json',
            'imgs': ['1bea552d-5295-4ee1-a43c-16b527a59859_0.jpg']
        },
        '6': {
            'json': 'bbeead62-b737-48e4-9807-99c2953b974b.json',
            'imgs': ['bbeead62-b737-48e4-9807-99c2953b974b_0.jpg']
        }

    }
    test = '3'

    past(os.path.join('../data/image/requests', case_tests[test]['json']),file_id=case_tests[test]['json'].replace('.json', ''))

    data = pickle.load(open('data.pkl', 'rb'))

    builder = DocumentBuilder()
    document = builder(**data)
    # print(document.info)
    # print(document.render())
    document.update_word_content_index()

    print(document.get_text_in_bbox(0, (940, 534, 1634, 740), keep_layout=True))
    sub_page = document.sub_page_in_bbox(0, (940, 534, 1634, 740), contain=True)

    images = []
    for img_name in case_tests[test]['imgs']:
        img = cv2.imread(
            os.path.join('../data/image/requests', img_name))
        images.append(img)
    document.show(images)
    print(document.export())
    # export_docs(document, 'data.docx')
